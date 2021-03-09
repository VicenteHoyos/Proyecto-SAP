from win32com import client
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches
import os
import getpass
import time

class Runnable:  
    def run(self, ses, code, row, utils):
        # must be redefined
        return            

class Logger:
    def __init__(self, id):
        # initialize document
        self.doc = Document()        
        self.paragraphRun = self.doc.add_paragraph().add_run()
        # initialize log's directory name
        self.path = os.path.expanduser('~'+getpass.getuser()) + '\\sapgui_scripts_logs\\log_'+id+'_'+str(datetime.now()).replace(" ","_").replace(":","_").replace(".","_")
        if not os.path.exists(self.path):
            os.makedirs(self.path)
        # initialize log entries counter
        self.num = 0
        #initialize id
        self.id = id
        
    def add(self, ses):
        # add data into document
        self.paragraphRun.add_text(ses.FindById("wnd[0]/sbar").text)
        self.num += 1
        filename = self.path+"\\Screenshot_"+str(self.num)
        ses.FindById("wnd[0]").HardCopy(filename)
        self.paragraphRun.add_picture(filename+'.bmp', width=Inches(5.0), height=Inches(2.5))

    def save(self):
        print(self.path+'\\doc_'+self.id+'.docx')
        self.doc.save(self.path+'\\doc_'+self.id+'.docx')
        
class Utils:
    def __init__(self, id):
        self.log = Logger(id)

class Transaction:
    
    def __init__(self, tcode, pathToData):
        self.tcode = tcode
        self.path = pathToData
        self.sapGui = client.GetObject("SAPGUI").GetScriptingEngine
        self.ses = self.sapGui.FindById("ses[0]")
        self.utils = Utils(tcode)
        
    def runScript(self, runnable):
        with open(self.path) as csvfile:                    
            reader = csv.DictReader(csvfile, delimiter = ';')
            for row in reader:
                try:
                    runnable.run(self.ses, self.tcode, row, self.utils)
                except:
                    time.sleep(2)
                    self.utils.log.add(self.ses)
        self.utils.log.save()            
             
    
